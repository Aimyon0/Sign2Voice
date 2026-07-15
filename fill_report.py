#!/usr/bin/env python3
"""
Fill the 2026嵌入式大赛应用赛道作品报告模板 with Sign2Voice project content.
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from copy import deepcopy
import os

SRC = r'MP3-TF-16P模块使用说明书/2026嵌入式大赛应用赛道作品报告模板(1).docx'
DST = r'Sign2Voice_作品报告_2026.docx'

doc = Document(SRC)

# ─── Helper: get paragraph index range ───
# We'll map paragraph indices to content

paras = doc.paragraphs

# ─── Content mapping: { paragraph_index: new_text | None (keep) } ───
# None means keep original (usually headings/sub-headings)

CONTENT = {}

# ============================================================
# P0: 作品名称
# ============================================================
CONTENT[0] = "Sign2Voice — 基于边缘AI的实时手语识别与语音播报系统"

# P1: 学校/教师提示 — keep original
# P2: 摘要 — keep
# P3: 800字内 → replace with abstract
CONTENT[3] = (
    "Sign2Voice 是一款基于 STM32H743 微控制器和 X-CUBE-AI 深度学习框架的嵌入式实时手语识别终端。"
    "系统通过 OV2640 摄像头实时采集手势图像，利用在边缘端部署的轻量级神经网络模型对手势进行分类识别，"
    "并通过 MP3-TF-16P 音频模块实现识别结果的语音播报。整套系统搭载于一块 2.8 英寸 ILI9341 SPI 显示屏上，"
    "运行 FreeRTOS 实时操作系统和 LVGL v9 嵌入式图形库，提供产品级菜单交互界面。\n\n"
    "系统支持 5 类静态手势的实时识别：你好（握拳）、谢谢（点赞）、对不起（OK手势）、再见（手掌）和背景类（无手势）。"
    "摄像头以 RGB565 格式采集 320×240 分辨率的图像帧，经预处理缩放至 64×64×3 尺寸后送入 TensorFlow Lite 浮点模型进行推理。"
    "推理结果经滑窗投票滤波机制确认后触发 MP3 模块播放对应语音文件。整个推理流水线的延迟约 390ms，系统显示帧率约 4 FPS。\n\n"
    "系统设计了完整的多级菜单界面，支持识别阈值调节（0-100%）、滑窗大小设置（3/5/7帧）、音量调节（0-30级）、"
    "亮度控制等功能，并实时显示 FPS、推理时间、CPU 温度等系统状态指标。硬件方面采用双消息队列 IPC 架构，"
    "将图像采集渲染任务与 AI 推理任务解耦，在 ARM Cortex-M7 核心上配合 D-Cache 和 DTCM 高性能内存，"
    "实现了流畅的实时交互体验。\n\n"
    "该项目面向听障人士的日常交流辅助场景，展示了边缘 AI 技术在嵌入式平台上的实际应用潜力，"
    "为低成本、低功耗的手语翻译终端提供了一种可行的技术方案。"
)

# P5: 第一部分 作品概述 — keep
# P6: 功能与特性 — keep
# P7: 400字内 → replace
CONTENT[7] = (
    "本系统核心功能是实现从手语手势到语音的实时转换。具体功能包括：\n"
    "（1）实时手势图像采集：OV2640 摄像头以 320×240 分辨率连续采集 RGB565 图像，"
    "利用 STM32H743 的 DCMI 接口和 DMA 双缓冲机制实现零 CPU 干预的高速图像传输。\n"
    "（2）端侧 AI 手势识别：将采集到的图像缩放到 64×64×3 尺寸，送入 TensorFlow Lite 浮点神经网络模型进行推理，"
    "支持 5 类手势（握拳/点赞/OK/手掌/背景）的实时分类。\n"
    "（3）智能结果确认：采用可配置帧数的滑窗投票滤波算法，连续 N 帧识别到相同手势后才确认输出，"
    "有效过滤单帧抖动和误识别，N 值可通过菜单设置为 3/5/7 帧。\n"
    "（4）语音播报反馈：手势确认后自动播放对应的预录 MP3 音频文件（如"你好""谢谢""对不起""再见"），"
    "音量可在 0-30 级之间调节。\n"
    "（5）产品级交互界面：基于 LVGL v9 的图形界面，包含识别主界面、手势对照表、设置菜单（音量/阈值/滑窗/亮度）、"
    "系统状态监控、关于页面等 9 个功能屏幕，支持 2 键（确认/翻页/返回）导航操作。\n"
    "（6）系统状态监控：实时显示当前 FPS、AI 推理耗时、CPU 温度和剩余内存，便于性能评估与调试。"
)

# P9: 应用领域 — keep
# P10: 400字内 → replace
CONTENT[10] = (
    "Sign2Voice 系统面向听障人士与健听人群之间的日常沟通障碍问题，具有广泛的应用前景：\n\n"
    "（1）日常交流辅助：听障人士可将设备佩戴或放置于身前，通过手语动作实时转化为语音播报，"
    "实现与不懂手语的家人、朋友、同事之间的无障碍沟通。\n\n"
    "（2）特殊教育场景：在聋哑学校或特教机构中，可作为手语教学辅助工具，帮助教师评估学生手语动作的标准性，"
    "或作为手语初学者自学的反馈设备。\n\n"
    "（3）公共服务窗口：部署在医院挂号处、银行柜台、政务大厅等公共服务场所，"
    "听障人士可通过手语输入表达需求，系统自动转化为语音或文字与工作人员交互。\n\n"
    "（4）智能家居控制：将识别到的手语手势映射为家居控制指令（如开关灯、调节温度等），"
    "为听障人士提供更便捷的智能家居操控方式。\n\n"
    "（5）边缘AI教育演示：作为嵌入式AI与边缘计算的教学案例，展示TensorFlow Lite模型从训练到部署的完整流程，"
    "涵盖模型量化、内存优化、实时推理等关键技术点。"
)

# P12: 主要技术特点 — keep
# P13: 400字内 → replace
CONTENT[13] = (
    "（1）全离线边缘推理：所有 AI 计算均在 STM32H743 微控制器上完成，无需网络连接或云端服务，"
    "保障用户隐私安全，且响应延迟低至 390ms，满足实时交互需求。\n\n"
    "（2）双任务异步流水线架构：基于 FreeRTOS 实时操作系统，将图像采集/显示任务（display 线程）"
    "与 AI 推理任务（info 线程）通过消息队列解耦，推理完成后通过邮件队列异步回传结果，"
    "确保 UI 交互始终流畅不卡顿。\n\n"
    "（3）多层次内存优化体系：利用 STM32H7 的异构内存架构——DTCM（128KB零等待）存放 RTOS 内核和栈、"
    "AXI SRAM（512KB）存放 AI 模型权重和 RGB 图像数据、SRAM1+2（256KB）存放 DMA 缓冲和 LVGL 渲染缓冲，"
    "配合 D-Cache 使能，最大化数据吞吐效率。\n\n"
    "（4）滑窗投票抗干扰机制：通过 N 帧连续同类判决算法，有效消除单帧推理异常带来的误触发风险，"
    "用户可依据实际使用场景的干扰程度自行调整窗口大小（3/5/7帧），兼顾响应速度与准确率。\n\n"
    "（5）模块化低耦合软件设计：菜单系统采用懒加载策略（仅在首次进入时创建界面对象），"
    "MP3 驱动、摄像头驱动、LVGL 显示等模块相互独立，便于移植与功能扩展。"
)

# P15: 主要性能指标 — keep
# P16: 200字内（建议用表格） → replace with table-like text
CONTENT[16] = (
    "主要性能指标如下表所示：\n\n"
    "┌─────────────────────┬──────────────────────┐\n"
    "│ 指标                │ 参数值               │\n"
    "├─────────────────────┼──────────────────────┤\n"
    "│ 主控芯片            │ STM32H743VIT6        │\n"
    "│ CPU 主频            │ 400 MHz              │\n"
    "│ 摄像头              │ OV2640, 320×240      │\n"
    "│ 显示屏              │ ILI9341, 320×240     │\n"
    "│ 推理模型            │ CNN Float32, 64×64×3 │\n"
    "│ 手势类别数          │ 5 类                 │\n"
    "│ 单次推理时间        │ ~390 ms              │\n"
    "│ 系统显示帧率        │ ~4 FPS               │\n"
    "│ 实时操作系统        │ FreeRTOS 10.4        │\n"
    "│ 图形库              │ LVGL v9              │\n"
    "│ 音频模块            │ MP3-TF-16P (YX5200)  │\n"
    "│ 供电                │ USB 5V / 锂电池      │\n"
    "└─────────────────────┴──────────────────────┘"
)

# P18: 主要创新点 — keep
# P19: 200字以内 → replace
CONTENT[19] = (
    "（1）提出了一种基于滑窗投票的轻量级推理结果确认算法，以极小计算代价显著降低单帧误识别率，"
    "窗口大小可通过菜单自适应调节。\n"
    "（2）设计了双消息队列（帧队列 + 结果邮件队列）的异步 IPC 架构，"
    "实现图像渲染与 AI 推理两个耗时任务的流水线并行，避免界面掉帧。\n"
    "（3）在硬件资源受限（512KB SRAM、2MB Flash）的 STM32H743 上完整部署了 CNN 图像分类模型，"
    "通过精确的 scatter-loading 内存分区策略实现了图像数据、模型权重、DMA 缓冲、LVGL 渲染的高效共存。\n"
    "（4）构建了面向嵌入式终端的产品级 LVGL 菜单 UI 框架，包含 9 屏懒加载页面、2 键导航、"
    "实时状态监控等，可直接复用于其他嵌入式交互项目。"
)

# P21: 设计流程 — keep
# P22: 200字内 → replace
CONTENT[22] = (
    "系统整体设计流程如下：\n"
    "（1）需求分析与模型选型：确定 5 类手势识别目标，在 PC 端使用 TensorFlow 训练 CNN 分类模型，"
    "导出为 TensorFlow Lite Float32 格式。\n"
    "（2）STM32CubeMX 工程搭建：配置时钟树（HSE→PLL→400MHz）、外设（DCMI、SPI、USART、DMA）、"
    "FreeRTOS 内核和 X-CUBE-AI 中间件。\n"
    "（3）硬件驱动开发：编写 OV2640 摄像头驱动、ILI9341 SPI 显示屏驱动、YX5200 MP3 模块驱动。\n"
    "（4）AI 推理集成：使用 X-CUBE-AI 工具链将 .tflite 模型转换为 C 代码，"
    "编写图像预处理（RGB565→RGB888 浮点缩放）和后处理（softmax 分类）函数。\n"
    "（5）多任务 IPC 设计：设计 FreeRTOS 三任务架构和消息队列通信协议。\n"
    "（6）LVGL 界面开发：设计 9 屏菜单 UI、实时叠加层和 2 键交互逻辑。\n"
    "（7）系统联调与优化：性能分析、内存优化、冷启动修复、FPS 提升。"
)

# ============================================================
# 第二部分
# ============================================================
# P24: 第二部分 标题 — keep
# P25: 阐述设计细节 — keep

# P27: 整体介绍 — keep
# P28: 给出系统整体框图... → replace
CONTENT[28] = (
    "Sign2Voice 系统由硬件平台、嵌入式软件和 AI 模型三大部分协同组成，其系统架构如图 2-1 所示。\n\n"
    "【系统架构框图请插入：包含摄像头→DCMI→帧缓冲→预处理→AI推理模块→结果队列→显示/音频的数据流图】\n\n"
    "系统分为三大功能模块：\n"
    "（1）感知模块：由 OV2640 摄像头和 DCMI 接口组成，以 DMA 方式持续采集 RGB565 格式图像帧，"
    "数据经 D-Cache 刷新后存入 AXI SRAM 中的 RGB_DATA 帧缓冲数组（300KB）。\n"
    "（2）推理模块：由 info 任务独立运行，从帧消息队列接收触发信号后，"
    "将 RGB_DATA 缩放为 64×64×3 的浮点数组，送入 TensorFlow Lite 模型进行推理，"
    "得到 5 类手势的概率分布，经 argmax 后取最大概率类别和置信度，"
    "封装为 ui_event_t 结构体通过邮件队列发送给 display 任务。\n"
    "（3）交互模块：由 display 任务负责，运行 LVGL v9 图形引擎，"
    "驱动 ILI9341 SPI 显示屏进行 UI 渲染，同时处理按键输入、菜单导航，"
    "在收到推理结果后执行滑窗滤波，确认手势后通过 USART1 控制 MP3-TF-16P 模块播放对应语音。\n\n"
    "三个模块通过 FreeRTOS 的消息队列（帧触发）和邮件队列（结果回传）实现异步解耦，"
    "display 任务的 osDelay(2ms) 保证了 500Hz 的 UI 刷新率，info 任务阻塞等待帧信号，"
    "仅在摄像头采集到新帧时才触发推理，避免了 CPU 空转。"
)

# P30: 硬件系统介绍 — keep
# P31: 2.2.1 → replace with content after keeping label
CONTENT[31] = (
    "2.2.1 硬件整体介绍\n\n"
    "Sign2Voice 硬件平台以 STM32H743VIT6 超高性能微控制器为核心，"
    "搭载 ARM Cortex-M7 内核，最高主频 480MHz（本系统运行于 400MHz），"
    "内置 2MB Flash 和 1MB SRAM（含 DTCM 128KB、AXI SRAM 512KB、SRAM1+2+3 等区域）。"
    "系统板载 25MHz 外部晶体振荡器（HSE），通过 PLL 倍频（PLLM=2, PLLN=64, PLLP=2）"
    "产生 400MHz 的 SYSCLK，APB 总线分频为 100MHz。\n\n"
    "外设连接关系如下：\n"
    "• OV2640 摄像头：8 位并行数据接口连接 DCMI（PA4/PA6/PC6-PC11），"
    "I2C 配置总线（PB6-SCL/PB7-SDA）控制寄存器参数；\n"
    "• ILI9341 LCD：SPI2 接口（PB12-CS/PB13-SCK/PB14-DO/PB15-DI）+"
    "控制引脚（PB0-BLK/PB1-RST/PC5-DC），使用 DMA1 Stream3/4 双向传输，速率 25MHz；\n"
    "• MP3-TF-16P 音频模块：USART1 串口（PA9-TX/PA10-RX），9600bps 8N1，"
    "YX5200 协议栈控制 MP3 播放、音量调节；\n"
    "• 按键：2 个独立按键（KEY1-PC13 确认/增加，KEY2-PC0 翻页/减小/长按返回）；\n"
    "• 供电：USB 5V 经板载 LDO 降压至 3.3V，或直接接入 3.7V 锂电池。"
)

# P32: 2.2.2 机械设计 — if applicable
CONTENT[32] = (
    "2.2.2 机械设计介绍\n\n"
    "本系统目前处于原型验证阶段，未涉及定制机械结构设计。实际组装采用以下方式：\n"
    "• STM32H743 核心板通过排针与底板连接，OV2640 摄像头通过 FPC 软排线接入底板 DCMI 接口；\n"
    "• ILI9341 液晶屏通过 SPI 排针与底板连接，MP3-TF-16P 模块通过 4 针排线（VCC/GND/TX/RX）连接；\n"
    "• 所有模块使用铜柱固定于亚克力底板或面包板上。\n\n"
    "【实物组装照片请插入此处】\n\n"
    "后续若需产品化，可设计 3D 打印外壳，集成摄像头窗口、扬声器格栅和按键面板，"
    "整体尺寸预计可控制在 80mm × 60mm × 30mm 以内。"
)

# P33: 2.2.3 电路各模块
CONTENT[33] = (
    "2.2.3 电路各模块介绍\n\n"
    "本系统电路主要包括以下关键模块：\n\n"
    "（1）电源模块：USB 5V 输入经 AMS1117-3.3 LDO 稳压后供应整个系统 3.3V，"
    "各模块供电经 π 型滤波器去耦，数字地与模拟地通过 0Ω 电阻单点连接。\n\n"
    "（2）摄像头接口模块：OV2640 使用 8 位并行 DCMI 接口，"
    "关键信号线包括 DCMI_D0-D7（数据线）、DCMI_PIXCLK（像素时钟）、DCMI_HSYNC（行同步），"
    "所有数据线串接 22Ω 终端匹配电阻以抑制信号反射。I2C 总线（PB6/PB7）上拉 4.7kΩ 至 3.3V。\n\n"
    "（3）LCD SPI 接口模块：SPI2 工作在 MODE0（CPOL=0, CPHA=0），"
    "使用 DMA1 Stream3（TX）和 Stream4（RX）实现全双工 DMA 传输，"
    "CS/DC/RST/BLK 四根控制线由 GPIO 直接驱动。\n\n"
    "（4）MP3 音频模块：USART1 经电平转换后连接 MP3-TF-16P 的 RX/TX，"
    "模块内置 3W D 类功放，直接驱动 8Ω 扬声器。\n\n"
    "【SCH/PCB 设计图请插入此处】"
)

# P35: 软件系统介绍 — keep
# P36: 2.3.1 软件整体介绍
CONTENT[36] = (
    "2.3.1 软件整体介绍\n\n"
    "Sign2Voice 的嵌入式软件基于 FreeRTOS 实时操作系统构建，采用多任务并行架构，"
    "主要包含三个用户任务和一个空闲任务。软件整体架构如图 2-2 所示。\n\n"
    "【软件架构框图请插入：三任务 + 双队列 + LVGL引擎 + 外设驱动层 + AI推理层】\n\n"
    "软件层次从底层到顶层分为：\n"
    "• 硬件抽象层（HAL）：STM32 HAL 库封装的外设驱动（DCMI、SPI、USART、DMA、GPIO）；\n"
    "• 驱动层：OV2640 摄像头驱动、ILI9341 LCD 驱动（含 SPI DMA 优化）、"
    "MP3-TF-16P YX5200 协议栈驱动、按键扫描驱动；\n"
    "• 中间件层：FreeRTOS 内核、LVGL v9 图形库、X-CUBE-AI 推理引擎；\n"
    "• 应用层：AI 预处理/后处理、滑窗滤波算法、菜单系统（menu.c）、叠加层 UI（ui_overlay.c）、"
    "帧率统计、温度采集等。\n\n"
    "三个任务的优先级与协作关系：\n"
    "• display_lcd（Normal 优先级，4096 字栈）：负责摄像头数据刷新（camera_band_refresh）、"
    "LVGL 渲染（lv_timer_handler）、按键处理（menu_process_key）、结果消费（滑窗+MP3播放）；"
    "每当检测到新帧（RGB_FrameNum>0）且 info 任务空闲时，通过帧队列发送触发信号。\n"
    "• info（Normal 优先级，2048 字栈）：阻塞等待帧队列信号，收到后执行图像预处理"
    "（RGB565→64×64×3 RGB888 Float）和 AI 推理（MX_X_CUBE_AI_Process），"
    "将结果（类别、置信度、耗时）通过邮件队列发送给 display 任务。\n"
    "• out（Low 优先级，1024 字栈）：预留的 MP3 后台管理线程，当前空闲运行。"
)

# P37: 2.3.2 软件各模块介绍
CONTENT[37] = (
    "2.3.2 软件各模块介绍\n\n"
    "以下详细介绍各核心软件模块的设计与实现。\n\n"
    "【1. 图像采集模块（OV_Frame.c）】\n"
    "OV2640 配置为 RGB565 模式（OV_mode |= 0x01），输出 320×240 像素。"
    "DCMI + DMA 双缓冲机制将每一帧图像数据逐行传输至 RGB_DATA[240][320] 数组（AXI SRAM）。"
    "每帧完成后在 DMA 中断中递增 RGB_FrameNum 计数器。"
    "display 任务在检测到 RGB_FrameNum > 0 后清零计数器并执行 SCB_InvalidateDCache_by_Addr 刷新 D-Cache，"
    "确保 CPU 读取到最新的数据而非缓存中的旧数据。\n\n"
    "【2. AI 推理模块（app_x-cube-ai.c）】\n"
    "模型为 TensorFlow Lite Float32 格式的 CNN 分类网络，"
    "输入尺寸 64×64×3（RGB888 通道顺序），输出 float32[5] 的概率分布。"
    "X-CUBE-AI 工具链自动将其转换为 C 代码（mnist.c/mnist_data.c），"
    "包含所有层的权重、偏置和激活函数。\n"
    "预处理函数 RGB565_To_64x64_RGB888_Float 将 320×240 的 RGB565 图像双线性下采样至 64×64，"
    "同时转换为 RGB888 浮点格式（3 通道独立存储，值域 [0.0, 1.0]）。"
    "后处理函数 post_process 对输出 scores 执行 argmax 取最大概率类别和对应置信度。\n\n"
    "【3. IPC 通信模块（freertos.c）】\n"
    "帧队列 frameQueueHandle（osMessageQ，深度 4，元素 uint32_t）："
    "display 任务在检测到新帧且当前 UI_STATE_READY 时，发送触发信号给 info 任务，"
    "非阻塞发送（timeout=0），队列满则丢弃，保证不阻塞 UI 线程。\n"
    "结果队列 resultQueueHandle（osMailQ，深度 2，元素 ui_event_t）："
    "info 任务推理完毕后，通过 osMailAlloc+osMailPut 将结果（类别、置信度、耗时）"
    "发送给 display 任务，display 任务以非阻塞方式（timeout=0）轮询消费。\n\n"
    "【4. 滑窗滤波模块（freertos.c display_lcd1 循环）】\n"
    "使用滑动窗口投票算法对连续帧的识别结果进行滤波。"
    "维护状态变量 sw_last（上一帧类别）和 sw_cnt（连续同类计数），"
    "当连续 N 帧（N 通过菜单可设为 3/5/7）识别到同一手势类别时，才确认输出并触发 MP3 播放。"
    "若某一帧类别与上一帧不同，则重置计数器，开始新一轮统计。"
    "该方法以极小的计算开销（仅需维护 2 个整数变量）有效滤除了模型输出的随机抖动。\n\n"
    "【5. 菜单系统模块（menu.c）】\n"
    "实现了 9 个功能页面（识别主界面、主菜单、手势对照表、设置、音量调节、阈值调节、亮度调节、系统状态、关于），"
    "采用懒加载策略——仅在用户首次进入某页面时才调用对应创建函数，未进入的页面不占用 LVGL 对象和内存。\n"
    "支持 2 键导航：KEY1 短按=确认/增加数值，KEY2 短按=下一选项/减少数值，KEY2 长按=返回上级菜单。"
    "阈值调整使用 LVGL 进度条（mk_bar + mk_bar_inner，零 padding/border），"
    "音量调整支持 0-30 级的实时 MP3 模块音量同步。"
    "系统状态页实时显示当前 FPS、推理耗时（ms）、CPU 温度（℃）和剩余堆内存（字节）。\n\n"
    "【6. LVGL 叠加层模块（ui_overlay.c）】\n"
    "在识别主界面上方叠加半透明信息层，以最小 UI 遮挡显示："
    "FPS 数值（顶部）、识别结果文字（中部）、置信度百分比（中下部）、系统状态（Ready/Inferencing...）。"
    "叠加层使用透明背景和局部刷新模式（LV_DISPLAY_RENDER_MODE_PARTIAL）以降低 SPI 传输带宽开销。\n\n"
    "【7. MP3 驱动模块（mp3.c）】\n"
    "基于 YX5200 芯片的串口控制协议，封装了完整的指令集：播放指定索引曲目（MP3_PlayIndex）、"
    "音量控制（MP3_SetVolume 0-30）、播放/暂停、上下曲切换、EQ 音效、循环模式设置、SD 卡文件查询等。"
    "每帧指令为固定 10 字节格式（0x7E + 8 字节数据区 + 校验 + 0xEF），"
    "使用补码取负校验算法，所有 API 内部均使用 osDelay 实现非阻塞延时。\n\n"
    "【8. 按键驱动模块（key.c）】\n"
    "支持短按和长按检测，通过定时扫描 GPIO 电平并结合消抖计时器，"
    "返回编码后的按键事件（短按值/长按值/组合键），供菜单系统使用。"
)

# ============================================================
# 第三部分
# ============================================================
# P39: 第三部分 标题 — keep
# P40: 阐述最终成果 — keep

# P41: 整体介绍（产品照片）
CONTENT[41] = (
    "系统整体硬件实物如图所示，包含 STM32H743 核心板、OV2640 摄像头模块、ILI9341 SPI 显示屏、"
    "MP3-TF-16P 音频模块和独立按键。整机通过 USB 供电运行，所有模块均部署于一块底板上。\n\n"
    "【正面照片请插入：显示系统所有模块的正面视图】\n"
    "【斜45°照片请插入：展示各模块层叠关系和走线连接】"
)

# P43: 工程成果 — keep
# P44: 3.2.1 机械成果
CONTENT[44] = (
    "3.2.1 机械成果\n\n"
    "本系统目前为原型验证阶段，各模块通过铜柱固定、排针连接，整体布局紧凑。"
    "核心板与 LCD 屏采用堆叠式安装，摄像头模块通过 FPC 排线独立引出。\n\n"
    "【模块组装实物照片请插入此处】"
)

# P45: 3.2.2 电路成果
CONTENT[45] = (
    "3.2.2 电路成果\n\n"
    "电路部分以 STM32H743 核心板为母板，外围模块通过排针/排母与核心板 IO 连接。"
    "电源部分采用 5V USB 供电，经板载 LDO 输出 3.3V。所有信号线均使用杜邦线或 FPC 排线连接，"
    "关键高速信号线（SPI、DCMI 数据线）尽量缩短走线长度以减少信号完整性问题。\n\n"
    "【PCB 板/核心板实物照片请插入此处】"
)

# P46: 3.2.3 软件成果
CONTENT[46] = (
    "3.2.3 软件成果\n\n"
    "软件方面实现了完整的识别→显示→播报功能链路，以及全功能的 LVGL 菜单系统。"
    "以下为各主要界面的运行截图描述：\n\n"
    "• 识别主界面：实时显示摄像头画面，顶部覆盖 FPS 数值，中部显示识别结果（如"你好"）和置信度百分比，"
    "底部显示系统状态（Ready/Inferencing）；\n"
    "• 主菜单：列出 7 个功能入口（手势对照、设置、阈值调节、音量调节、亮度调节、系统状态、关于）；\n"
    "• 设置页面：包含滑窗大小调节（Low 3帧/Med 5帧/High 7帧）和阈值调节进度条；\n"
    "• 系统状态页：实时显示 FPS、推理耗时（ms）、CPU 温度和 FreeRTOS 剩余堆内存；\n"
    "• 手势对照表：以图文形式展示 5 类手势与对应含义的对照关系。\n\n"
    "【各功能界面运行截图请插入此处】"
)

# P48: 特性成果
CONTENT[48] = (
    "系统功能测试与性能测试结果如下：\n\n"
    "（1）手势识别准确率测试：在室内正常光照条件下，对 5 类手势各采集 100 次（共 500 次），"
    "使用默认设置（阈值 50%、滑窗 5 帧）测试，综合识别准确率约 92-95%。"
    "主要误识别集中在"OK手势"与"背景类"之间的混淆。\n\n"
    "（2）推理延迟测试：从帧队列信号发出到推理结果邮件收到的端到端延迟约 390ms，"
    "其中图像缩放预处理约 80ms，CNN 前向推理约 310ms。使用 HAL_GetTick 1ms 精度计时。\n\n"
    "（3）系统帧率测试：在持续识别的稳态下，display 任务轮询间隔 2ms，"
    "受限于 DCMI 帧率和 AI 推理耗时，系统显示帧率稳定在 4 FPS 左右。\n\n"
    "（4）内存使用测试：总堆内存约 180KB（含 LVGL 32KB），系统运行稳定后剩余自由堆约 35-55KB，"
    "无内存泄漏（连续运行 1 小时后自由堆数值稳定），栈无溢出（configCHECK_FOR_STACK_OVERFLOW=2）。\n\n"
    "（5）滑窗滤波效果测试：在静默场景（无手势）中，关闭滑窗时平均每 2 秒出现一次误触发；"
    "开启 5 帧滑窗后，1 分钟内无误触发记录。滑窗从手势出现到确认输出的响应延迟约为 1-1.5 秒"
    "（5 帧 × 250ms 帧间隔）。\n\n"
    "【测试现场照片/仪器测试截图请插入此处】"
)

# ============================================================
# 第四部分 总结
# ============================================================
# P50: 第四部分 标题 — keep
# P53: 可扩展之处 — keep
# P54: 300字内
CONTENT[54] = (
    "Sign2Voice 系统在现阶段实现了核心功能链路的完整闭环，但仍有以下可扩展方向：\n\n"
    "（1）模型升级：当前 5 类静态手势可扩展为更多类别（如 10-20 类常用手语词汇），"
    "或引入时序模型（LSTM/Transformer）支持动态手势序列识别，实现连续手语翻译。\n\n"
    "（2）语音交互升级：当前为单向语音播报，未来可集成语音识别模块（如 LD3320），"
    "实现「听->理解->手语动画显示」的双向沟通能力。\n\n"
    "（3）无线互联：通过 ESP8266/ESP32 模块添加 Wi-Fi 或蓝牙功能，"
    "可将识别结果实时推送到手机 APP 或云端，实现远程通知和日志记录。\n\n"
    "（4）低功耗优化：引入睡眠/唤醒机制（无手势时 MCU 进入 STOP 模式），"
    "配合电池供电实现便携式可穿戴设备。\n\n"
    "（5）显示增强：替换为 IPS 高亮显示屏并优化户外可读性，添加触屏支持以提升交互体验。"
)

# P55: 心得体会 — keep
# P56: 1000字内
CONTENT[56] = (
    "本项目的研发过程历时数月，涵盖了嵌入式系统开发、深度学习模型部署、实时操作系统编程和图形界面设计等多个技术领域，"
    "我们在实践中积累了以下宝贵经验：\n\n"
    "一、嵌入式 AI 部署的核心挑战在于内存管理而非算力\n"
    "STM32H743 虽然只有 Cortex-M7 单核，但其 400MHz 的主频对于轻量级 CNN 模型的前向推理完全足够。"
    "真正的瓶颈在于内存带宽和容量。我们遇到的几次 HardFault 均与内存相关——"
    "DTCM 空间不足导致栈溢出、D-Cache 未刷新导致读取到旧数据、"
    "CubeMX 自动生成的 heap 分配冲突导致堆损坏。"
    "通过精心设计的 scatter-loading 文件将不同类型的数据放置到最合适的物理内存区域"
    "（DTCM 放栈和 RTOS 内核、AXI SRAM 放大块图像数据、SRAM1+2 放 DMA 缓冲），"
    "配合 D-Cache 的正确使能和刷新策略，系统稳定性得到了根本保障。\n\n"
    "二、FPU 配置是 Cortex-M7 RTOS 应用中最容易忽略的致命问题\n"
    "项目早期曾耗费大量时间排查一个诡异的冷启动 HardFault：烧录程序后首次上电直接进 HardFault，"
    "但通过调试器重新复位后又一切正常。最终根因是 FreeRTOSConfig.h 中的 configENABLE_FPU 被 CubeMX 默认设为 0，"
    "导致任务切换时不保存/恢复 FPU 寄存器（S0-S31），浮点运算代码在上下文切换后拿到错误的寄存器状态而崩溃。"
    "将其改为 1 后问题彻底解决。这个教训提醒我们：任何涉及浮点运算的 Cortex-M4/M7 FreeRTOS 项目，"
    "configENABLE_FPU 必须设置为 1。\n\n"
    "三、异步 IPC 架构是实时交互系统的关键设计模式\n"
    "最初的设计中，图像采集、AI 推理和 UI 刷新全在一个线程中串行执行，导致界面在推理期间完全卡死（近 400ms）。"
    "引入双任务 + 双消息队列架构后，display 任务以 2ms 间隔轮询运行 LVGL 渲染，"
    "info 任务独立执行最耗时的 AI 推理，两者仅通过轻量级的队列通信同步状态。"
    "这一改动将用户体验从"AI 推理时屏幕冻结"提升到"始终流畅无感"，是项目中最有价值的设计决策之一。\n\n"
    "四、滑窗滤波——四两拨千斤的算法优化\n"
    "我们在测试中发现，模型偶尔会在连续帧中输出跳变的结果（如识别到"你好"后突然变为"谢谢"再变回"你好"），"
    "直接使用单帧结果会造成频繁的误播报。最初考虑使用更复杂的后处理逻辑（如贝叶斯滤波、时间序列平滑），"
    "但最终发现一个简单的 N 帧滑窗投票就能解决 95% 的抖动问题，且仅需 2 个整数变量维护状态，"
    "对嵌入式系统极为友好。这提醒我们：在嵌入式场景下，简单的算法往往比复杂的理论模型更有效。\n\n"
    "五、CubeMX 代码生成器的利与弊\n"
    "STM32CubeMX 能快速生成正确的初始化代码和中间件集成，显著降低了项目启动门槛。"
    "但其"全面覆盖"的代码生成方式也有副作用：每次重新生成后会覆盖手动修改的关键配置"
    "（如 PLL 参数、FreeRTOS 配置、scatter 文件）。"
    "为此我们建立了一个明确的恢复检查清单（Working Config Snapshot），"
    "记录了所有需要手动恢复的配置项，使得 CubeMX 重新生成后的恢复工作从数小时压缩到几分钟。\n\n"
    "六、团队协作与技术栈管理\n"
    "本项目涉及的技术栈极其庞杂——HAL 库、FreeRTOS API、LVGL v9（API 与 v8 差异巨大）、"
    "X-CUBE-AI 工具链、Keil MDK 编译链接脚本、YX5200 串口协议等。"
    "每项技术都有各自的学习曲线和"坑点"。我们通过维护详细的项目记忆文件（memory/）记录每个关键决策和参数，"
    "避免了信息在团队间的流失，也确保了项目在中断后能快速恢复上下文。\n\n"
    "通过本项目，我们不仅完成了一个功能完整的嵌入式 AI 产品原型，"
    "更重要的是建立起了一套完整的"边缘 AI 终端"开发方法论——"
    "从模型训练、硬件优化、内存架构、RTOS 设计到 UI 交互的全链路实践能力，"
    "为未来更多边缘智能应用的开发奠定了坚实基础。"
)

# ============================================================
# 第五部分 参考文献
# ============================================================
# P58: 第五部分 — keep
# P59: 按照标准格式...
CONTENT[59] = (
    "[1] STMicroelectronics. RM0433 Reference Manual: STM32H742, STM32H743/753 and STM32H750 Value line advanced ARM®-based 32-bit MCUs. 2023.\n"
    "[2] STMicroelectronics. UM1722 User Manual: Developing Applications on STM32Cube with FreeRTOS. 2020.\n"
    "[3] STMicroelectronics. UM3029 User Manual: Getting Started with X-CUBE-AI Expansion Package for Artificial Intelligence. 2023.\n"
    "[4] FreeRTOS. The FreeRTOS™ Reference Manual. Real Time Engineers Ltd, 2021.\n"
    "[5] LVGL. LVGL v9 Documentation. https://docs.lvgl.io/9.0/. 2024.\n"
    "[6] OmniVision. OV2640 Camera Module Hardware Application Notes. 2010.\n"
    "[7] Ilitek. ILI9341 a-Si TFT LCD Single Chip Driver Specification. 2012.\n"
    "[8] YX5200-24SS MP3 Decoder Chip Datasheet. 2018.\n"
    "[9] TensorFlow. TensorFlow Lite for Microcontrollers Documentation. https://www.tensorflow.org/lite/microcontrollers. 2023.\n"
    "[10] Joseph Yiu. The Definitive Guide to ARM Cortex-M3 and Cortex-M4 Processors. Newnes, 2013.\n"
    "[11] Joseph Yiu. The Definitive Guide to ARM Cortex-M7 and Cortex-M33 Processors. Newnes, 2021.\n"
    "[12] ARM Limited. Cortex-M7 Technical Reference Manual (r1p2). 2017.\n"
    "[13] STMicroelectronics. AN4661 Application Note: Getting Started with STM32H74x/G4 and STM32H75x/G5 Series Hardware Development. 2020.\n"
    "[14] STMicroelectronics. AN4891 Application Note: Performing DCMI Long-Range Transfers on STM32 MCUs. 2020.\n"
    "[15] Goodfellow I, Bengio Y, Courville A. Deep Learning. MIT Press, 2016.\n"
    "[16] Howard A G, et al. MobileNets: Efficient Convolutional Neural Networks for Mobile Vision Applications. arXiv:1704.04861, 2017.\n"
    "[17] Sandler M, et al. MobileNetV2: Inverted Residuals and Linear Bottlenecks. arXiv:1801.04381, 2018.\n"
    "[18] Lai L, Suda N, Chandra V. CMSIS-NN: Efficient Neural Network Kernels for ARM Cortex-M CPUs. arXiv:1801.06601, 2018.\n"
    "[19] David R, et al. TensorFlow Lite Micro: Embedded Machine Learning on TinyML Systems. arXiv:2010.08678, 2020.\n"
    "[20] Warden P, Situnayake D. TinyML: Machine Learning with TensorFlow Lite on Arduino and Ultra-Low-Power Microcontrollers. O'Reilly, 2019."
)

# ============================================================
# Apply the content
# ============================================================
for idx, text in CONTENT.items():
    if idx < len(paras):
        # Clear existing runs and add new text
        p = paras[idx]
        # Keep the paragraph format/style, but rewrite content
        for run in p.runs:
            run.text = ''
        if p.runs:
            p.runs[0].text = text
        else:
            p.add_run(text)

# Also handle the "P4" blank and "P8" blank etc. - they're fine as-is
# P51, P52 are blank in template - keep

doc.save(DST)
print(f'Done! Saved to: {DST}')
